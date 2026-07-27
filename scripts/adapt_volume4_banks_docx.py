#!/usr/bin/env python3
"""Prepare the beginner-friendly Volume 4 Part 1 central-banks DOCX."""

from __future__ import annotations

import argparse
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if " ".join(paragraph.text.split()) == " ".join(text.split()):
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def clear_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_run_with_format(paragraph: Paragraph, text: str, run_properties) -> None:
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, deepcopy(run_properties))


def replace_text(paragraph: Paragraph, text: str) -> None:
    run_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs else None
    clear_runs(paragraph)
    add_run_with_format(paragraph, text, run_properties)


def replace_labeled_text(paragraph: Paragraph, label: str, body: str) -> None:
    label_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs else None
    body_index = next(
        (index for index, run in enumerate(paragraph.runs) if not run.bold and run.text.strip()),
        len(paragraph.runs) - 1,
    )
    body_properties = (
        deepcopy(paragraph.runs[body_index]._r.rPr) if paragraph.runs else None
    )
    clear_runs(paragraph)
    add_run_with_format(paragraph, f"{label}  ", label_properties)
    add_run_with_format(paragraph, body, body_properties)


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def insert_labeled_paragraph_before(
    target: Paragraph,
    formatting_source: Paragraph,
    label: str,
    body: str,
) -> Paragraph:
    element = OxmlElement("w:p")
    target._p.addprevious(element)
    paragraph = Paragraph(element, target._parent)
    if formatting_source._p.pPr is not None:
        paragraph._p.insert(0, deepcopy(formatting_source._p.pPr))
    label_properties = (
        deepcopy(formatting_source.runs[0]._r.rPr) if formatting_source.runs else None
    )
    body_index = next(
        (
            index
            for index, run in enumerate(formatting_source.runs)
            if not run.bold and run.text.strip()
        ),
        len(formatting_source.runs) - 1,
    )
    body_properties = (
        deepcopy(formatting_source.runs[body_index]._r.rPr)
        if formatting_source.runs
        else None
    )
    add_run_with_format(paragraph, f"{label}  ", label_properties)
    add_run_with_format(paragraph, body, body_properties)
    return paragraph


def restart_numbered_list(document: Document, paragraphs: list[Paragraph]) -> None:
    """Assign a fresh numbering instance so a new pedagogical list starts at 1."""
    first_num_properties = paragraphs[0]._p.pPr.find(qn("w:numPr"))
    original_num_id = first_num_properties.find(qn("w:numId")).get(qn("w:val"))
    numbering = document.part.numbering_part.element
    original_num = numbering.find(f".//{{{numbering.nsmap['w']}}}num[@{{{numbering.nsmap['w']}}}numId='{original_num_id}']")
    abstract_num_id = original_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_num_ids = [int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))]
    new_num_id = str(max(existing_num_ids, default=0) + 1)

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), new_num_id)
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), abstract_num_id)
    new_num.append(abstract_reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    new_num.append(level_override)
    numbering.append(new_num)

    for paragraph in paragraphs:
        num_properties = paragraph._p.pPr.find(qn("w:numPr"))
        num_properties.find(qn("w:numId")).set(qn("w:val"), new_num_id)


def adapt_document(source: Path, output: Path) -> None:
    document = Document(source)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                run.text = run.text.replace("VOLUME 6", "VOLUME 4")

    replacements = {
        "VOLUME 4 · NIVEAU DÉBUTANT · ÉDITION 2026": (
            "VOLUME 4 · PARTIE 1 · NIVEAU DÉBUTANT · ÉDITION 2026"
        ),
        (
            "Ce volume prolonge les cours précédents. Le cours macroéconomique a expliqué "
            "l'inflation, l'emploi et la croissance ; le volume géopolitique a montré comment "
            "un choc traverse l'économie ; les volumes techniques ont appris à lire le prix et "
            "le volume. Ici, tu vas comprendre l'acteur qui relie ces éléments : la banque centrale."
        ): (
            "Cette Partie 1 ouvre le Volume 4. Les volumes précédents ont posé les bases de "
            "l'investissement, du risque et de la lecture du marché. Ici, tu vas comprendre "
            "l'institution qui fixe l'orientation monétaire. La Partie 2 appliquera ensuite ce "
            "cadre aux données d'inflation, d'emploi, de croissance et de consommation."
        ),
        "Ce que tu sauras faire à la fin": "Objectifs de la Partie 1",
        (
            "Comprendre les taux directeurs, les opérations de marché, les réserves, le QE, le "
            "QT, la forward guidance et le contrôle de la courbe des taux."
        ): (
            "Comprendre d'abord les taux directeurs, les réserves et les opérations de marché, "
            "puis situer le QE, le QT et le YCC comme outils d'approfondissement."
        ),
        "Plan": "Parcours de la Partie 1",
        "Du taux directeur au bilan : ce que la banque change réellement.": (
            "Du taux directeur au bilan : l'essentiel d'abord, les outils exceptionnels ensuite."
        ),
        "3.4 Quantitative easing — QE": "3.4 Approfondissement — Quantitative easing (QE)",
        "3.5 Quantitative tightening — QT": "3.5 Approfondissement — Quantitative tightening (QT)",
        "3.7 Contrôle de la courbe des taux — YCC": (
            "3.7 Approfondissement — Contrôle de la courbe des taux (YCC)"
        ),
        "3.8 Prêteur en dernier ressort": "3.8 Approfondissement — Prêteur en dernier ressort",
        "3.9 Interventions de change": "3.9 Approfondissement — Interventions de change",
    }
    for old_text, new_text in replacements.items():
        replace_text(find_paragraph(document, old_text), new_text)

    plan_items = [
        find_paragraph(document, "Rôle et importance des banques centrales."),
        find_paragraph(document, "Organisation d'une décision de politique monétaire."),
        find_paragraph(document, "Outils conventionnels et non conventionnels."),
        find_paragraph(document, "Transmission vers l'économie réelle."),
        find_paragraph(document, "Réactions des différents marchés."),
        find_paragraph(document, "Deux cas réels : Fed en août 2022 et Banque du Japon en décembre 2022."),
        find_paragraph(document, "Méthode d'analyse, erreurs fréquentes et glossaire."),
    ]
    restart_numbered_list(document, plan_items)

    chapter_titles = {
        1: ("Qu'est-ce qu'une banque centrale ?", "Qu'est-ce qu'une banque centrale ?"),
        2: ("Comment une décision est-elle prise ?", "Comment une décision est-elle prise ?"),
        3: ("Les outils de politique monétaire", "Les outils de politique monétaire"),
        4: ("La transmission vers l'économie", "La transmission vers l'économie"),
        5: ("L'impact sur les marchés financiers", "L'impact sur les marchés financiers"),
        6: ("Cas 1 — Fed et Jackson Hole 2022", "Cas 1 — Fed et Jackson Hole 2022"),
        7: ("Cas 2 — Banque du Japon et YCC", "Cas 2 — Banque du Japon et YCC"),
        8: ("Méthode d'analyse pour débutant", "Méthode d'analyse d'une décision monétaire"),
    }
    for chapter_number, (original_title, adapted_title) in chapter_titles.items():
        title = find_paragraph(document, original_title)
        replace_text(title, f"{chapter_number}. {adapted_title}")
        title.style = document.styles["Heading 1"]
        remove_paragraph(find_paragraph(document, f"CHAPITRE {chapter_number}"))

    first_tool_heading = find_paragraph(document, "3.1 Les taux directeurs")
    central_idea = find_paragraph(
        document,
        (
            "IDÉE CENTRALE Le marché ne réagit pas seulement à une hausse ou à une baisse de "
            "taux. Il compare la décision, le communiqué, les prévisions et le discours à ce "
            "qu'il avait déjà anticipé."
        ),
    )
    insert_labeled_paragraph_before(
        first_tool_heading,
        central_idea,
        "PARCOURS DÉBUTANT",
        (
            "Pour une première lecture, maîtrise les sections 3.1 à 3.3 et 3.6. Les sections "
            "consacrées au QE, au QT, au YCC, au prêteur en dernier ressort et aux interventions "
            "de change sont des approfondissements : elles peuvent être relues dans un second temps."
        ),
    )

    bridge = find_paragraph(
        document,
        (
            "LIEN AVEC LES VOLUMES PRÉCÉDENTS La macroéconomie explique pourquoi la banque agit ; "
            "la géopolitique décrit les chocs auxquels elle répond ; l'analyse technique montre "
            "comment le marché matérialise la nouvelle ; la gestion du risque protège lorsque la "
            "réaction diffère du scénario."
        ),
    )
    replace_labeled_text(
        bridge,
        "PASSAGE À LA PARTIE 2",
        (
            "Tu connais maintenant le mandat, les outils et les canaux de transmission des banques "
            "centrales. La Partie 2 ne répétera pas ces mécanismes : elle montrera comment chaque "
            "publication d'inflation, d'emploi, de croissance ou de consommation peut modifier les "
            "anticipations de politique monétaire et les scénarios de marché."
        ),
    )

    document.core_properties.title = "Volume 4 · Partie 1 — Les banques centrales"
    document.core_properties.subject = (
        "Comprendre le rôle, les décisions, les outils et la transmission des banques centrales"
    )
    document.core_properties.keywords = (
        "banques centrales, politique monétaire, taux directeurs, QE, QT, YCC, marchés financiers"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--copy", type=Path)
    args = parser.parse_args()
    adapt_document(args.source, args.output)
    if args.copy:
        args.copy.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.output, args.copy)


if __name__ == "__main__":
    main()
