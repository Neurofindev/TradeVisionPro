#!/usr/bin/env python3
"""Adapt the geopolitical course as Volume 4, Part 3."""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def normalized(text: str) -> str:
    return " ".join(text.split())


def find_paragraph(document: Document, text: str) -> Paragraph:
    target = normalized(text)
    for paragraph in document.paragraphs:
        if normalized(paragraph.text) == target:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def clear_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def replace_text(paragraph: Paragraph, text: str) -> None:
    run_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs else None
    clear_runs(paragraph)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def replace_labeled_text(paragraph: Paragraph, label: str, body: str) -> None:
    label_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs else None
    body_run = next(
        (run for run in paragraph.runs if not run.bold and run.text.strip()),
        paragraph.runs[-1] if paragraph.runs else None,
    )
    body_properties = deepcopy(body_run._r.rPr) if body_run is not None else None
    clear_runs(paragraph)
    label_run = paragraph.add_run(f"{label}  ")
    if label_properties is not None:
        label_run._r.insert(0, label_properties)
    body_output = paragraph.add_run(body)
    if body_properties is not None:
        body_output._r.insert(0, body_properties)


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def update_header_volume(document: Document) -> None:
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                run.text = run.text.replace("VOLUME 5", "VOLUME 4")


def adapt_document(source: Path, output: Path) -> None:
    document = Document(source)
    update_header_volume(document)

    replacements = {
        "VOLUME 5 · NIVEAU DÉBUTANT · ÉDITION 2026": (
            "VOLUME 4 · PARTIE 3 · NIVEAU DÉBUTANT · ÉDITION 2026"
        ),
        "Dix graphiques TradingView en bougies japonaises · cas historiques sourcés": (
            "Dix graphiques TradingView et un schéma pédagogique · cas historiques sourcés"
        ),
        "Ce que tu sauras faire à la fin": "Objectifs de la Partie 3",
        "Plan du cours": "Parcours de la Partie 3",
        "La chaîne de transmission vers les marchés": (
            "Appliquer la chaîne de transmission aux chocs géopolitiques"
        ),
        "Lire chaque grande classe d'actifs": (
            "Lire l'effet géopolitique sur chaque classe d'actifs"
        ),
        (
            "Un choc de demande détruit surtout la croissance. Un choc d'offre réduit la "
            "production disponible et augmente les coûts. Le second est plus difficile pour "
            "une banque centrale : soutenir l'activité peut entretenir l'inflation, tandis "
            "que combattre l'inflation peut affaiblir davantage la croissance."
        ): (
            "La Partie 2 a présenté les régimes de croissance et d'inflation. Ici, applique "
            "ce cadre au choc géopolitique : un choc de demande pèse surtout sur l'activité, "
            "tandis qu'un choc d'offre réduit la production disponible et augmente les coûts. "
            "Ce dernier peut placer la banque centrale devant un arbitrage plus difficile."
        ),
        (
            "Le rendement obligataire et la valorisation des actions dépendent des attentes "
            "de taux, d'inflation et de sécurité. Lors d'un choc déflationniste, les obligations "
            "d'État solides peuvent être recherchées et leurs rendements baisser. Lors d'un choc "
            "énergétique inflationniste, les rendements peuvent au contraire monter. Il n'existe "
            "donc pas de réaction obligataire universelle à la géopolitique."
        ): (
            "La Partie 1 a expliqué comment les anticipations de taux atteignent les obligations "
            "et les actions. Dans un scénario géopolitique, détermine maintenant quelle force "
            "domine : recherche de sécurité, inflation importée, ralentissement de la croissance "
            "ou hausse du risque souverain. Il n'existe pas de réaction obligataire universelle."
        ),
        (
            "Le prix révèle la différence entre le scénario attendu et l'information nouvelle. "
            "Une mauvaise nouvelle peut faire monter le marché si elle est moins grave que prévu. "
            "Une bonne nouvelle peut provoquer une baisse si les investisseurs étaient déjà "
            "positionnés pour un résultat parfait."
        ): (
            "Le Volume 3 a montré comment le prix et le volume matérialisent les anticipations. "
            "Ici, utilise-les pour vérifier le scénario géopolitique : une mauvaise nouvelle peut "
            "faire monter le marché si elle est moins grave que prévu, tandis qu'une nouvelle "
            "positive peut décevoir si un résultat parfait était déjà intégré."
        ),
        (
            "Une devise est toujours un prix relatif. Pour lire EUR/USD, GBP/USD ou USD/CNH, "
            "il faut comparer deux économies. Une monnaie peut baisser parce que son pays est "
            "directement exposé, parce que la banque centrale devrait assouplir sa politique, "
            "parce que les capitaux sortent ou parce que les importations d'énergie deviennent "
            "plus coûteuses."
        ): (
            "Rappel de la Partie 2 : une devise est un prix relatif entre deux économies. Pour "
            "un choc géopolitique, compare leur exposition directe, la réponse probable de leurs "
            "banques centrales, les flux de capitaux et le coût des importations d'énergie."
        ),
        (
            "Les obligations sont le marché le plus ambigu. Une peur de récession peut faire "
            "baisser les rendements grâce à la recherche de sécurité. Mais un choc d'offre "
            "énergétique peut renforcer l'inflation et faire monter les rendements. Il faut donc "
            "déterminer quelle force domine : fuite vers la qualité, inflation ou risque souverain."
        ): (
            "Applique ici le cadre obligataire de la Partie 1 : une peur de récession peut faire "
            "baisser les rendements par recherche de sécurité, alors qu'un choc d'offre "
            "énergétique peut renforcer l'inflation et les faire monter. Identifie la force "
            "dominante : fuite vers la qualité, inflation ou risque souverain."
        ),
        (
            "CE QUE CELA SIGNIFIE Le marché a réévalué en quelques heures la valeur relative "
            "de la livre et le risque attaché au Royaume-Uni. L'ouverture, la clôture, la mèche "
            "et le volume relient ce dossier aux Volumes 3 et 4."
        ): (
            "CE QUE CELA SIGNIFIE Le marché a réévalué en quelques heures la valeur relative "
            "de la livre et le risque attaché au Royaume-Uni. L'ouverture, la clôture, la mèche "
            "et le volume relient ce dossier au Volume 3, consacré à la lecture technique du marché."
        ),
        "11.4 Lire TradingView après l'analyse fondamentale": (
            "11.4 Lire TradingView après l'analyse géopolitique"
        ),
        "12.3 Ce que les graphiques de ce volume enseignent": (
            "12.3 Ce que les graphiques de cette Partie 3 enseignent"
        ),
        "Conclusion générale — relier les cinq volumes": (
            "Conclusion générale — relier les quatre volumes"
        ),
        (
            "Le Volume 1 apprend à comprendre un actif et à séparer valeur, comportement du prix "
            "et gestion du risque. Le Volume 2 montre qu'un mécanisme caché peut finir par dominer "
            "le graphique. Le Volume 3 donne des outils pour lire tendance, momentum et participation. "
            "Le Volume 4 revient au langage des bougies. Le Volume 5 ajoute les rapports de puissance "
            "capables de modifier les flux, l'inflation, la croissance et les primes de risque."
        ): (
            "Le Volume 1 apprend à comprendre un actif et à distinguer l'analyse d'entreprise, "
            "le comportement du prix et la gestion du risque. Le Volume 2 montre, à travers des "
            "cas historiques, comment exposition, financement, gouvernance et liquidité transforment "
            "une thèse. Le Volume 3 fournit les outils de lecture technique. Dans le Volume 4, les "
            "banques centrales forment la Partie 1, les données macroéconomiques la Partie 2 et les "
            "rapports de puissance capables de modifier les flux, les coûts et les primes de risque "
            "la Partie 3."
        ),
        (
            "Le cours macroéconomique se place au centre de cette chaîne : il mesure les conséquences "
            "qui finissent par apparaître dans le CPI, le PCE, l'emploi, le PIB, les ventes au détail "
            "et les décisions des banques centrales. La géopolitique ne remplace donc aucune analyse "
            "précédente. Elle fournit un scénario fondamental supplémentaire que la macroéconomie et "
            "le marché doivent confirmer."
        ): (
            "La Partie 2 mesure les conséquences qui peuvent finir par apparaître dans le CPI, le PCE, "
            "l'emploi, le PIB, les ventes au détail et les décisions des banques centrales. La "
            "géopolitique ne répète ni ne remplace ces analyses : elle ajoute un scénario fondamental "
            "dont le canal économique, les données et le comportement du marché doivent confirmer la portée."
        ),
    }
    for old_text, new_text in replacements.items():
        replace_text(find_paragraph(document, old_text), new_text)

    thread = find_paragraph(
        document,
        (
            "FIL CONDUCTEUR La géopolitique construit le scénario. La macroéconomie mesure ses "
            "conséquences. Le graphique montre comment le marché les intègre. La gestion du risque "
            "protège le capital lorsque le scénario reste incertain."
        ),
    )
    replace_labeled_text(
        thread,
        "POSITION DANS LE VOLUME 4",
        (
            "La Partie 1 a expliqué les banques centrales et la transmission monétaire ; la Partie 2, "
            "les publications macroéconomiques. Cette Partie 3 ajoute l'origine géopolitique de certains "
            "chocs et suit leur passage vers les flux réels, les coûts, les anticipations et les prix."
        ),
    )
    objectives = find_paragraph(document, "Objectifs de la Partie 3")
    objectives._p.addnext(thread._p)

    chapter_titles = [
        "Comprendre la géopolitique",
        "Acteurs, ressources et instruments",
        "Appliquer la chaîne de transmission aux chocs géopolitiques",
        "Lire l'effet géopolitique sur chaque classe d'actifs",
        "Dossier 1 — Brexit",
        "Dossier 2 — Abqaiq",
        "Dossier 3 — Guerre commerciale États-Unis–Chine",
        "Dossier 4 — Invasion de l'Ukraine",
        "Dossier 5 — Semi-conducteurs",
        "Dossier 6 — Mer Rouge et routes maritimes",
        "Méthode d'analyse et de décision",
        "Limites, erreurs et discipline",
    ]
    for number, title_text in enumerate(chapter_titles, start=1):
        title = find_paragraph(document, title_text)
        replace_text(title, f"{number}. {title_text}")
        title.style = document.styles["Heading 1"]
        remove_paragraph(find_paragraph(document, f"CHAPITRE {number}"))

    document.core_properties.title = (
        "Volume 4 · Partie 3 — Géopolitique et marchés financiers"
    )
    document.core_properties.subject = (
        "Comprendre comment les chocs géopolitiques se transmettent à l'économie et aux marchés"
    )
    document.core_properties.keywords = (
        "géopolitique, géoéconomie, sanctions, conflits, matières premières, devises, marchés financiers"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    adapt_document(args.source, args.output)


if __name__ == "__main__":
    main()
